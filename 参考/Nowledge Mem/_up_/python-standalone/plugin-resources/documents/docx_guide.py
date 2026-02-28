#!/usr/bin/env python3
"""
DOCX Generation Knowledge Base
Generated: 2025-11-29
Purpose: Document best practices and patterns for python-docx usage

Run this script to see comprehensive python-docx documentation.
The agent should run this when uncertain about DOCX generation.
"""

def print_docx_knowledge():
    """Print comprehensive DOCX generation knowledge and best practices"""

    knowledge = """
# 📝 PYTHON-DOCX GENERATION GUIDE

## 1. CORE CONCEPTS & ARCHITECTURE

### What is python-docx?
- Low-level XML manipulation library for Word documents (.docx)
- Based on OpenXML (OOXML) format (same as PPTX)
- Designed for content creation and document structure
- Perfect for generating reports, letters, contracts, technical docs

### Key Classes
```python
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Main objects:
# - Document: Root container
# - Paragraph: Text block (highest level in document flow)
# - Run: Inline text with consistent formatting
# - Table: Tabular data structure
# - Section: Page-level settings
```

## 2. DOCUMENT PHILOSOPHY

### ✅ STRENGTHS
- Excellent for structured documents (reports, proposals)
- Built-in styles system
- Easy paragraph/section management
- Native table support with rich styling
- Simple formatting (fonts, colors, alignment)

### ⚠️ LIMITATIONS
- Page layout is complex (margins, headers, footers require XML manipulation)
- Limited design capabilities
- No automatic layout/positioning like PowerPoint
- Styling is less flexible than HTML/CSS
- Template support is minimal (can load, but not as clean as PPTX)

### Design Approach (RECOMMENDED)
**Focus on content and structure, not design:**
1. Write clean, well-structured content
2. Use built-in Word styles
3. Add tables, lists, images for data presentation
4. Save and let Word format the rest

## 3. WORKFLOW PATTERNS

### Pattern 1: From Scratch (Most Common)
```python
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Add title
title = doc.add_heading('Report Title', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add content
p = doc.add_paragraph('This is the introduction paragraph.')
p.paragraph_format.space_after = Pt(12)

# Add subsection
doc.add_heading('Section 1', level=1)
doc.add_paragraph('Content here...')

# Save
doc.save('output.docx')
```

### Pattern 2: Modify Existing Document
```python
from docx import Document

# Load existing document
doc = Document('template.docx')

# Add new content
doc.add_heading('New Section', 1)
doc.add_paragraph('New content added programmatically')

# Save with new name
doc.save('modified.docx')
```

### Pattern 3: Content Population with Data
```python
from docx import Document

doc = Document()

# Add structured content with data
for item in data_list:
    doc.add_heading(item['title'], level=2)
    doc.add_paragraph(item['description'])

    # Add table for metrics
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    for metric, value in item['metrics'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)

    doc.add_paragraph()  # Add spacing

doc.save('output.docx')
```

## 4. TEXT & PARAGRAPH FORMATTING

### Heading Levels
```python
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

doc.add_heading('Title (Level 0)', 0)           # Largest
doc.add_heading('Heading (Level 1)', 1)         # Sections
doc.add_heading('Subheading (Level 2)', 2)      # Subsections
doc.add_heading('Sub-subheading (Level 3)', 3)  # Details
doc.add_paragraph('Normal paragraph')            # Body text
```

### Run-level Formatting (Inline)
```python
p = doc.add_paragraph()

# Add text with different formatting
p.add_run('Bold text').bold = True
p.add_run(' - ')
p.add_run('Italic text').italic = True
p.add_run(' - ')
r = p.add_run('Colored text')
r.font.color.rgb = RGBColor(192, 0, 0)  # Red

# Font properties
r.font.size = Pt(12)
r.font.name = 'Calibri'
```

### Paragraph Formatting
```python
from docx.enum.text import WD_ALIGN_PARAGRAPH

p = doc.add_paragraph('Centered paragraph')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # LEFT, RIGHT, CENTER, JUSTIFY

# Spacing
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.line_spacing = 1.5

# Indentation
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.first_line_indent = Inches(0.25)
```

### Bullet Points & Lists
```python
doc = Document()

doc.add_paragraph('Item 1', style='List Bullet')
doc.add_paragraph('Item 2', style='List Bullet')
doc.add_paragraph('Sub-item', style='List Bullet 2')

doc.add_paragraph('First item', style='List Number')
doc.add_paragraph('Second item', style='List Number')
doc.add_paragraph('Sub-item', style='List Number 2')
```

## 5. COLOR SYSTEM

### RGB Colors
```python
from docx.shared import RGBColor

# Define color constants
TITLE_COLOR = RGBColor(31, 78, 121)      # Dark blue
ACCENT_COLOR = RGBColor(192, 0, 0)       # Red
TEXT_COLOR = RGBColor(89, 89, 89)        # Gray
SUCCESS_COLOR = RGBColor(46, 204, 113)   # Green
WARNING_COLOR = RGBColor(230, 126, 34)   # Orange

# Apply to text
run = p.add_run('Colored text')
run.font.color.rgb = ACCENT_COLOR
```

### Paragraph Borders & Shading
```python
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Add shading (background color)
def shade_paragraph(paragraph, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    paragraph._element.get_or_add_pPr().append(shading_elm)

p = doc.add_paragraph('Highlighted paragraph')
shade_paragraph(p, 'FFFF00')  # Yellow background
```

## 6. TABLES

### Table Creation
```python
from docx.shared import Inches

# Create table: rows, cols
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Access rows and cells
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Header 1'
hdr_cells[1].text = 'Header 2'
hdr_cells[2].text = 'Header 3'

# Fill data
data = [
    ('Row 1 Col 1', 'Row 1 Col 2', 'Row 1 Col 3'),
    ('Row 2 Col 1', 'Row 2 Col 2', 'Row 2 Col 3'),
    ('Row 3 Col 1', 'Row 3 Col 2', 'Row 3 Col 3'),
]

for i, (col1, col2, col3) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = col1
    row_cells[1].text = col2
    row_cells[2].text = col3
```

### Table Styling
```python
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

# Shade cell
def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Apply to header row
for cell in table.rows[0].cells:
    shade_cell(cell, 'C5D9F1')  # Light blue

    # Format header text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(12)
```

### Table Width & Alignment
```python
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
table.autofit = False  # Disable auto-fit
table.allow_autofit = False

for row in table.rows:
    row.cells[0].width = Inches(2)
    row.cells[1].width = Inches(3)
```

## 7. IMAGES

### Inserting Images
```python
from docx.shared import Inches

# Add image with specified width
doc.add_picture('/absolute/path/to/image.png', width=Inches(6))

# Add to specific paragraph
p = doc.add_paragraph()
run = p.add_run()
run.add_picture('/path/to/image.png', width=Inches(5))

# Image in table cell
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture('/path/to/image.png', width=Inches(2))
```

## 8. PAGE SETUP & SECTIONS

### Document Margins
```python
from docx.shared import Inches

doc = Document()
sections = doc.sections

for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
```

### Page Breaks
```python
# Break to new page
doc.add_page_break()
```

### Headers & Footers
```python
from docx.shared import Inches, Pt

section = doc.sections[0]

# Header
header = section.header
p = header.paragraphs[0]
p.text = 'Document Title - Header'

# Footer
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Page 1'
```

## 9. STYLES

### Built-in Styles
```python
# Available styles:
# - Normal
# - Heading 1, Heading 2, Heading 3, ...
# - List Bullet, List Number
# - Title
# - Subtitle
# - List Paragraph
# - Quote
# - Intense Quote

doc.add_heading('Title', style='Title')
doc.add_paragraph('Quote text', style='Quote')
doc.add_paragraph('Item 1', style='List Bullet')
```

### Custom Style Creation
```python
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# Create custom style
styles = doc.styles
style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(16)
style.font.bold = True
style.font.color.rgb = RGBColor(31, 78, 121)

# Use custom style
doc.add_paragraph('My custom heading', style='CustomHeading')
```

## 10. COMPLETE EXAMPLE: Report Generator

```python
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_report(title, sections_data, output_path):
    \"\"\"Generate a professional report document.\"\"\"

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Report Generated: 2025-11-29')
    doc.add_page_break()

    # Table of contents
    doc.add_heading('Table of Contents', level=1)
    for section in sections_data:
        doc.add_paragraph(section['name'], style='List Number')

    doc.add_page_break()

    # Content sections
    for section in sections_data:
        doc.add_heading(section['name'], level=1)
        doc.add_paragraph(section['content'])

        # Add table if available
        if 'table_data' in section and section['table_data']:
            headers = section['table_data'][0]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Headers
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)

            # Data rows
            for row_data in section['table_data'][1:]:
                row = table.add_row()
                for i, cell_data in enumerate(row_data):
                    row.cells[i].text = str(cell_data)

        doc.add_page_break()

    doc.save(output_path)
    print(f"Created: {output_path}")

# Usage
sections = [
    {
        'name': 'Executive Summary',
        'content': 'This report provides an overview of Q4 performance...',
    },
    {
        'name': 'Financial Results',
        'content': 'Revenue increased by 15% compared to Q3.',
        'table_data': [
            ['Metric', 'Q3', 'Q4', 'Change'],
            ['Revenue', '$1.2M', '$1.38M', '+15%'],
            ['Profit', '$300K', '$350K', '+17%'],
        ]
    }
]
generate_report("Q4 Report", sections, "q4_report.docx")
```

## 11. BEST PRACTICES

### DO ✅
- Use absolute file paths
- Define color constants once, reuse everywhere
- Use built-in styles when possible
- Add proper spacing between sections
- Test document by opening in Word
- Structure content logically (headings, paragraphs, tables)
- Use meaningful variable names
- Add comments for complex formatting

### DON'T ❌
- Try to create pixel-perfect layouts (Word is flow-based, not absolute positioning)
- Use relative paths
- Hardcode colors
- Skip margins/spacing (readability suffers)
- Assume all systems have same fonts (use standard fonts like Calibri, Arial)
- Mix concerns (content generation + styling should be separate)
- Forget to save document before exiting

## 12. QUICK REFERENCE

| Task | Code |
|------|------|
| Create document | `Document()` |
| Load document | `Document('file.docx')` |
| Add heading | `doc.add_heading('Text', level)` |
| Add paragraph | `doc.add_paragraph('Text')` |
| Bold text | `run.bold = True` |
| Colored text | `run.font.color.rgb = RGBColor(255, 0, 0)` |
| Add table | `doc.add_table(rows, cols)` |
| Add image | `doc.add_picture('path', width=Inches(6))` |
| Page break | `doc.add_page_break()` |
| Save | `doc.save('output.docx')` |
| Set margins | `section.top_margin = Inches(1)` |
| Align text | `p.alignment = WD_ALIGN_PARAGRAPH.CENTER` |

---

**Key Insight:** python-docx is for structured content, not pixel-perfect design.
Focus on clear content + good structure, and Word will handle the rest! 📄✨
"""

    print(knowledge)


if __name__ == '__main__':
    print_docx_knowledge()

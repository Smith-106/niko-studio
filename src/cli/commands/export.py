"""
export command - Export Functionality
======================================

Exports content to various formats (Markdown, JSON, DOCX).
"""

import click
import json
from pathlib import Path
from datetime import datetime
from rich.console import Console
from rich.panel import Panel


@click.command()
@click.argument("source", type=click.Path(exists=True))
@click.option(
    "--format", "-f",
    type=click.Choice(["md", "json", "docx", "txt", "html"]),
    default="md",
    help="Export format"
)
@click.option(
    "--output", "-o",
    type=click.Path(),
    help="Output file path (default: auto-generated)"
)
@click.option(
    "--template", "-t",
    type=click.Choice(["novel", "screenplay", "report", "plain"]),
    default="plain",
    help="Export template"
)
@click.option(
    "--include-meta",
    is_flag=True,
    help="Include metadata in export"
)
@click.pass_context
def export(
    ctx: click.Context,
    source: str,
    format: str,
    output: str,
    template: str,
    include_meta: bool
) -> None:
    """Export content to various formats.

    Supports exporting drafts, sessions, and evaluation results.

    Examples:
        niko export draft.json --format md
        niko export session.json --format docx
        niko export chapter.md --format html --template novel
    """
    console: Console = ctx.obj.get("console", Console())
    source_path = Path(source)

    console.print(f"\n[bold blue]Niko Studio - Export[/]")
    console.print(f"[dim]Source:[/] {source}")
    console.print(f"[dim]Format:[/] {format}")
    console.print(f"[dim]Template:[/] {template}")

    # Read source content
    content = source_path.read_text(encoding="utf-8")

    # Determine source type and parse
    if source_path.suffix == ".json":
        try:
            data = json.loads(content)
            is_session = "role" in str(data) or isinstance(data, list)
        except json.JSONDecodeError:
            data = {"content": content}
            is_session = False
    else:
        data = {"content": content}
        is_session = False

    # Generate output path if not specified
    if not output:
        output = source_path.with_suffix(f".{format}")
        # Avoid overwriting source
        if output == source_path:
            output = source_path.with_name(f"{source_path.stem}_export.{format}")

    output_path = Path(output)

    # Export based on format
    if format == "md":
        exported = _export_markdown(data, template, include_meta, is_session)
    elif format == "json":
        exported = _export_json(data, include_meta)
    elif format == "txt":
        exported = _export_text(data, is_session)
    elif format == "html":
        exported = _export_html(data, template, include_meta, is_session)
    elif format == "docx":
        exported = _export_docx(data, template, output_path, is_session)
        if exported is None:
            return  # DOCX writes directly
    else:
        console.print(f"[red]Unsupported format: {format}[/]")
        return

    # Write output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(exported, encoding="utf-8")

    console.print(Panel(
        f"[bold]Exported successfully![/]\n\n"
        f"Output: [cyan]{output_path}[/]\n"
        f"Size: {len(exported)} characters",
        title="[green]Export Complete[/]",
        border_style="green"
    ))


def _export_markdown(data: dict, template: str, include_meta: bool, is_session: bool) -> str:
    """Export to Markdown format."""
    lines = []

    if include_meta:
        lines.append("---")
        lines.append(f"exported: {datetime.now().isoformat()}")
        lines.append(f"template: {template}")
        lines.append("---\n")

    if is_session and isinstance(data, list):
        # Session export
        lines.append("# Chat Session\n")
        for msg in data:
            role = "**You:**" if msg.get("role") == "user" else "**Niko:**"
            lines.append(f"{role}\n")
            lines.append(f"{msg.get('content', '')}\n")
            lines.append("---\n")
    else:
        # Content export
        content = data.get("content", str(data))

        if template == "novel":
            lines.append("# Chapter\n")
            lines.append(content)
        elif template == "screenplay":
            lines.append("# SCREENPLAY\n")
            lines.append("```")
            lines.append(content)
            lines.append("```")
        else:
            lines.append(content)

    return "\n".join(lines)


def _export_json(data: dict, include_meta: bool) -> str:
    """Export to JSON format."""
    if include_meta:
        data = {
            "metadata": {
                "exported": datetime.now().isoformat(),
                "version": "1.0"
            },
            "data": data
        }
    return json.dumps(data, indent=2, ensure_ascii=False)


def _export_text(data: dict, is_session: bool) -> str:
    """Export to plain text format."""
    if is_session and isinstance(data, list):
        lines = []
        for msg in data:
            role = "You" if msg.get("role") == "user" else "Niko"
            lines.append(f"[{role}]")
            lines.append(msg.get("content", ""))
            lines.append("")
        return "\n".join(lines)
    else:
        return data.get("content", str(data))


def _export_html(data: dict, template: str, include_meta: bool, is_session: bool) -> str:
    """Export to HTML format."""
    content = data.get("content", str(data)) if isinstance(data, dict) else str(data)

    # Basic HTML template
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Niko Studio Export</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 2rem; line-height: 1.6; }}
        .user {{ color: #2563eb; }}
        .assistant {{ color: #059669; }}
        .meta {{ color: #6b7280; font-size: 0.875rem; }}
        hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 1.5rem 0; }}
    </style>
</head>
<body>
"""

    if include_meta:
        html += f'<p class="meta">Exported: {datetime.now().isoformat()}</p>\n'

    if is_session and isinstance(data, list):
        html += "<h1>Chat Session</h1>\n"
        for msg in data:
            role_class = "user" if msg.get("role") == "user" else "assistant"
            role_label = "You" if msg.get("role") == "user" else "Niko"
            html += f'<p class="{role_class}"><strong>{role_label}:</strong></p>\n'
            html += f'<p>{msg.get("content", "")}</p>\n'
            html += "<hr>\n"
    else:
        if template == "novel":
            html += "<h1>Chapter</h1>\n"
        html += f"<div>{content.replace(chr(10), '<br>')}</div>\n"

    html += """</body>
</html>"""

    return html


def _export_docx(data: dict, template: str, output_path: Path, is_session: bool) -> str:
    """Export to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        # Fallback: save as markdown with .docx extension note
        content = _export_markdown(data, template, True, is_session)
        note = "# Note: python-docx not installed. Content saved as Markdown.\n\n"
        return note + content

    doc = Document()

    # Add title
    doc.add_heading("Niko Studio Export", 0)

    content = data.get("content", str(data)) if isinstance(data, dict) else str(data)

    if is_session and isinstance(data, list):
        for msg in data:
            role = "You" if msg.get("role") == "user" else "Niko"
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(msg.get("content", ""))
            doc.add_paragraph()
    else:
        # Add content paragraphs
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Add metadata
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(f"Exported: {datetime.now().isoformat()}").italic = True

    # Save document
    doc.save(str(output_path))

    return None  # Indicates file was written directly
